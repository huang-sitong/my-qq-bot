"""文档导入核心服务。

策略：
- .txt / .json / .docx / .xlsx 走 LangChain 生态的轻量 loader；未安装
  langchain-community 时使用项目内轻量 fallback，保证脚本在最小依赖下可跑通。
- .pdf 解析做 3 重降级（见 .others/mineru.md）：
  ① MinerU 精准解析 API（v4，需 BOT_DOC_MINERU_API_KEY/ENDPOINT：签名上传 → 轮询 → 取 full.md）；
  ② MinerU Agent 轻量解析 API（v1，免 Token，≤10MB/≤20 页，BOT_DOC_MINERU_AGENT_ENABLED 开关）；
  ③ 本地 LangChain / pypdf 解析。
- 解析结果统一为 langchain_core.documents.Document，再按文档类型选择结构感知
  切分器（可用时优先 langchain-text-splitters，否则使用内置简单切分）。
- 入库到独立 documents collection，并做 file_hash 去重。
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import logging
from datetime import datetime
from pathlib import Path

from langchain_core.documents import Document

from bot.package.config import BotConfig
from bot.package.domain.repositories import DocumentRepository

from .document_store import DocumentStore
from .service import TS_FMT

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".txt", ".json"}

# ---------------------------------------------------------------------------
# 可选依赖探测
# ---------------------------------------------------------------------------


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Document(page_content=text, metadata={"source": str(path)})]


def _load_json(path: Path) -> list[Document]:
    data = json.loads(path.read_text(encoding="utf-8"))
    text = json.dumps(data, ensure_ascii=False, indent=2)
    return [Document(page_content=text, metadata={"source": str(path)})]


def _load_docx(path: Path) -> list[Document]:
    # LangChain 轻量 loader
    try:
        from langchain_community.document_loaders import Docx2txtLoader

        try:
            return Docx2txtLoader(str(path)).load()
        except Exception:
            logger.debug("LangChain Docx2txtLoader failed, fallback to python-docx", exc_info=True)
    except Exception:
        logger.debug("langchain_community unavailable, use python-docx fallback", exc_info=True)

    # python-docx fallback
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - depends on optional deps
        raise RuntimeError(
            "解析 .docx 需要安装 langchain-community + docx2txt 或 python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return [Document(page_content="\n".join(parts), metadata={"source": str(path)})]


def _load_xlsx(path: Path) -> list[Document]:
    # LangChain 轻量 loader
    try:
        from langchain_community.document_loaders import OpenpyxlLoader

        return OpenpyxlLoader(str(path)).load()
    except Exception:
        logger.debug("langchain_community OpenpyxlLoader unavailable, use openpyxl fallback", exc_info=True)

    # openpyxl fallback（项目已有依赖）
    import openpyxl

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for worksheet in workbook.worksheets:
        parts.append(f"# Sheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            cells = [str(cell) for cell in row if cell is not None]
            if cells:
                parts.append(" | ".join(cells))
    workbook.close()
    return [Document(page_content="\n".join(parts), metadata={"source": str(path)})]


def _load_pdf_mineru(path: Path, config: BotConfig) -> list[Document] | None:
    """第 1 重降级：MinerU 精准解析 API（v4，需 Token/Endpoint）；失败返回 None。"""
    from . import mineru_client

    if mineru_client.mineru_base_url(config) is None:
        # 未配置精准解析（无 endpoint 且无 API Key），继续走下一重降级
        return None

    markdown = mineru_client.parse_pdf(path, config)
    if not markdown or not markdown.strip():
        logger.warning("MinerU precise parse returned empty text for %s", path)
        return None
    return [
        Document(
            page_content=markdown,
            metadata={"source": str(path), "parser": "mineru", "format": "markdown"},
        )
    ]


def _load_pdf_mineru_agent(path: Path, config: BotConfig) -> list[Document] | None:
    """第 2 重降级：MinerU Agent 轻量解析 API（v1，免 Token）；失败返回 None。"""
    from . import mineru_client

    if not getattr(config, "document_mineru_agent_enabled", True):
        return None

    markdown = mineru_client.parse_pdf_agent(path, config)
    if not markdown or not markdown.strip():
        logger.warning("MinerU agent parse returned empty text for %s", path)
        return None
    return [
        Document(
            page_content=markdown,
            metadata={"source": str(path), "parser": "mineru-agent", "format": "markdown"},
        )
    ]


def _load_pdf_langchain(path: Path) -> list[Document]:
    try:
        from langchain_community.document_loaders import PyPDFLoader

        try:
            return PyPDFLoader(str(path)).load()
        except Exception:
            logger.debug("LangChain PyPDFLoader failed, fallback to pypdf", exc_info=True)
    except Exception:
        logger.debug("langchain_community PyPDFLoader unavailable, use pypdf fallback", exc_info=True)

    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - depends on optional deps
        raise RuntimeError(
            "解析 .pdf 需要配置 MinerU（BOT_DOC_MINERU_ENDPOINT/API_KEY）或安装 langchain-community + pypdf"
        ) from exc

    reader = PdfReader(str(path))
    docs: list[Document] = []
    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": str(path), "page": page_index + 1},
                )
            )
    return docs


def _load_pdf(path: Path, config: BotConfig) -> list[Document]:
    # 1. MinerU 精准解析（v4）
    docs = _load_pdf_mineru(path, config)
    if docs:
        return docs
    # 2. MinerU Agent 轻量解析（v1，免 Token）
    docs = _load_pdf_mineru_agent(path, config)
    if docs:
        return docs
    # 3. 本地 LangChain / pypdf
    docs = _load_pdf_langchain(path)
    for doc in docs:
        doc.metadata.setdefault("parser", "langchain")
        doc.metadata.setdefault("format", "text")
    return docs


_LOADERS = {
    ".txt": _load_text,
    ".json": _load_json,
    ".docx": _load_docx,
    ".xlsx": _load_xlsx,
    ".pdf": _load_pdf,
}


# ---------------------------------------------------------------------------
# 切分
# ---------------------------------------------------------------------------


def _simple_split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list[str]:
    """不带 langchain-text-splitters 时的内置简单切分。"""
    if chunk_size <= 0:
        return [text] if text else []
    if len(text) <= chunk_size:
        return [text] if text else []

    separators = ["\n\n", "\n", "。", "！", "？", "；", " "]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            split_at = -1
            for sep in separators:
                pos = text.rfind(sep, start + 1, end + 1)
                if pos > start:
                    split_at = pos
                    break
            if split_at > start:
                end = split_at
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= len(text):
            break
        start = max(end - chunk_overlap, start + 1)
    return chunks


def _split_documents(docs: list[Document], chunk_size: int, chunk_overlap: int) -> list[Document]:
    """把解析出的 Document 切成适合向量化/检索的 chunks。"""
    try:
        from langchain_text_splitters import (  # type: ignore[import-not-found]
            MarkdownHeaderTextSplitter,
            RecursiveCharacterTextSplitter,
        )
    except Exception:
        logger.debug("langchain_text_splitters unavailable, use built-in splitter", exc_info=True)
        MarkdownHeaderTextSplitter = None
        RecursiveCharacterTextSplitter = None

    result: list[Document] = []
    for doc in docs:
        text = doc.page_content or ""
        if not text.strip():
            continue

        base_meta = dict(doc.metadata or {})
        markdown = base_meta.get("format") == "markdown"

        if markdown and MarkdownHeaderTextSplitter is not None:
            splitter = MarkdownHeaderTextSplitter(
                headers_to_split_on=[
                    ("#", "H1"),
                    ("##", "H2"),
                    ("###", "H3"),
                    ("####", "H4"),
                ],
                strip_headers=False,
            )
            chunks = splitter.split_text(text)
            for chunk in chunks:
                merged = {**base_meta, **chunk.metadata}
                merged["page_content"] = chunk.page_content
                result.append(
                    Document(page_content=chunk.page_content, metadata=merged)
                )
            continue

        if RecursiveCharacterTextSplitter is not None:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            pieces = splitter.split_text(text)
        else:
            pieces = _simple_split_text(text, chunk_size, chunk_overlap)

        for piece in pieces:
            result.append(Document(page_content=piece, metadata=dict(base_meta)))
    return result


# ---------------------------------------------------------------------------
# 入口
# ---------------------------------------------------------------------------


def _file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _build_metadata(
    path: Path,
    chunk: Document,
    doc_id: str,
    file_hash: str,
    chunk_index: int,
    imported_at: str,
) -> dict:
    ext = path.suffix.lower()
    base = {
        "thread_id": "documents",
        "doc_id": doc_id,
        "file_hash": file_hash,
        "file_name": path.name,
        "file_type": ext.lstrip("."),
        "page": int(chunk.metadata.get("page", 0) or 0),
        "chunk_index": chunk_index,
        "source_path": str(path),
        "imported_at": imported_at,
        "content": chunk.page_content,
    }
    # 保留切分器/解析器带来的附加 metadata，但不允许覆盖上面结构化字段
    return {**chunk.metadata, **base}


async def ingest_files(
    config: BotConfig,
    paths: list[str | Path],
    store: DocumentRepository | None = None,
) -> list[dict]:
    """导入一批文档到知识库，返回每个文件的导入结果摘要。"""
    owns_store = store is None
    if owns_store:
        store = DocumentStore(config, collection=getattr(config, "document_collection", "documents"))
    imported: list[dict] = []
    try:
        for raw_path in paths:
            path = Path(raw_path)
            if not path.is_file():
                logger.warning("skip missing file: %s", path)
                imported.append({"path": str(path), "status": "missing", "chunks": 0})
                continue

            ext = path.suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                logger.warning("unsupported file type: %s", path)
                imported.append({"path": str(path), "status": "unsupported", "chunks": 0})
                continue

            file_hash = _file_hash(path)
            if await store.has_doc(file_hash):
                logger.info("skip duplicate file: %s (%s)", path, file_hash[:12])
                imported.append({"path": str(path), "status": "duplicate", "chunks": 0})
                continue

            try:
                loader = _LOADERS.get(ext)
                if loader is None:  # pragma: no cover - guarded by supported extensions
                    continue

                if ext == ".pdf":
                    docs = await asyncio.to_thread(loader, path, config)
                else:
                    docs = await asyncio.to_thread(loader, path)
                chunks = _split_documents(
                    docs,
                    chunk_size=getattr(config, "document_chunk_size", 500),
                    chunk_overlap=getattr(config, "document_chunk_overlap", 50),
                )
                if not chunks:
                    logger.warning("no text extracted from %s", path)
                    imported.append({"path": str(path), "status": "empty", "chunks": 0})
                    continue

                doc_id = f"doc-{file_hash}"
                imported_at = datetime.now().strftime(TS_FMT)
                texts = [c.page_content for c in chunks]
                metadatas = [
                    _build_metadata(path, c, doc_id, file_hash, i, imported_at)
                    for i, c in enumerate(chunks)
                ]
                await store.add_texts(texts, metadatas)
                logger.info("imported %s: %d chunks", path, len(chunks))
                imported.append({
                    "path": str(path),
                    "status": "imported",
                    "chunks": len(chunks),
                    "doc_id": doc_id,
                })
            except Exception:
                logger.exception("failed to import %s", path)
                imported.append({"path": str(path), "status": "error", "chunks": 0})
    finally:
        if owns_store and store is not None:
            store.close()
    return imported


__all__ = [
    "SUPPORTED_EXTENSIONS",
    "DocumentStore",
    "ingest_files",
]
