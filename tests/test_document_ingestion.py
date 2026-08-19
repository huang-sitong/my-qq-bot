"""文档导入/存储测试。

覆盖：
- 文本、JSON、Excel 的轻量解析；
- 内置简单切分器；
- ingest_files 的入库与 file_hash 去重流水线；
- DocumentStore 真实 milvus-lite 建集合与去重。
"""

import asyncio
import hashlib
import json

from langchain_core.documents import Document

from bot.package.config import BotConfig
from bot.package.knowledge.document_ingestion import (
    _load_docx,
    _load_json,
    _load_pdf_mineru,
    _load_pdf_mineru_agent,
    _load_text,
    _load_xlsx,
    _simple_split_text,
    ingest_files,
)
from bot.package.knowledge.document_store import DocumentStore
from bot.package.tools.builtin.search_documents import search_documents

TXT = """\
这是第一段。用于测试文本解析。

这是第二段，包含一些足够长的内容。我们需要验证切分器可以按段落和句子切分。
重复内容用于凑长度。重复内容用于凑长度。重复内容用于凑长度。
"""


class FakeDocumentStore:
    """模拟 DocumentStore，避免测试依赖真实 Milvus。"""

    def __init__(self) -> None:
        self.added: list[tuple[list[str], list[dict]]] = []
        self.hashes: set[str] = set()

    async def has_doc(self, file_hash: str) -> bool:
        return file_hash in self.hashes

    async def add_texts(self, texts: list[str], metadatas: list[dict]) -> None:
        self.added.append((texts, metadatas))
        if metadatas:
            self.hashes.add(metadatas[0]["file_hash"])

    def close(self) -> None:
        pass


class FakeEmbedder:
    async def embed_query(self, query):
        return [1.0, 0.0, 0.0, 0.0]

    async def embed_documents(self, contents):
        return [[1.0, 0.0, 0.0, 0.0] for _ in contents]

    def close(self):
        pass


def test_load_text(tmp_path):
    path = tmp_path / "a.txt"
    path.write_text("hello\nworld", encoding="utf-8")
    docs = _load_text(path)
    assert len(docs) == 1
    assert docs[0].page_content == "hello\nworld"


def test_load_json(tmp_path):
    path = tmp_path / "a.json"
    path.write_text(json.dumps({"name": "知识库", "tags": ["pdf", "docx"]}, ensure_ascii=False), encoding="utf-8")
    docs = _load_json(path)
    assert len(docs) == 1
    assert "知识库" in docs[0].page_content


def test_load_docx(tmp_path):
    from docx import Document

    path = tmp_path / "a.docx"
    document = Document()
    document.add_paragraph("这是 docx 测试内容")
    document.add_paragraph("第二段")
    document.save(path)

    docs = _load_docx(path)
    assert len(docs) == 1
    assert "docx 测试内容" in docs[0].page_content


def test_load_xlsx(tmp_path):
    import openpyxl

    path = tmp_path / "a.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["列A", "列B"])
    ws.append(["苹果", "香蕉"])
    wb.save(path)
    wb.close()

    docs = _load_xlsx(path)
    assert len(docs) == 1
    assert "Sheet1" in docs[0].page_content
    assert "苹果" in docs[0].page_content


def test_simple_split_text_respects_chunk_size():
    chunks = _simple_split_text("一二三四五六七八九十" * 30, chunk_size=20, chunk_overlap=0)
    assert len(chunks) > 1
    assert all(len(c) <= 30 for c in chunks)  # 受分隔符影响可能略超过 chunk_size
    assert "".join(chunks).replace(" ", "") == ("一二三四五六七八九十" * 30)


def test_ingest_files_continues_after_parse_error(tmp_path):
    config = BotConfig(embed_dimensions=4)
    store = FakeDocumentStore()
    bad = tmp_path / "bad.json"
    bad.write_text("this is not json", encoding="utf-8")
    good = tmp_path / "good.txt"
    good.write_text("正常的文本内容", encoding="utf-8")

    results = asyncio.run(ingest_files(config, [bad, good], store=store))
    by_path = {r["path"]: r["status"] for r in results}
    assert by_path[str(bad)] == "error"
    assert by_path[str(good)] == "imported"


def test_ingest_files_imports_and_deduplicates(tmp_path):
    config = BotConfig(
        embed_dimensions=4,
        document_chunk_size=200,
        document_chunk_overlap=20,
    )
    store = FakeDocumentStore()
    path = tmp_path / "note.txt"
    path.write_text(TXT, encoding="utf-8")

    results = asyncio.run(ingest_files(config, [path], store=store))
    assert results[0]["status"] == "imported"
    assert results[0]["chunks"] > 0
    assert store.added

    file_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    assert store.added[0][1][0]["file_hash"] == file_hash
    assert store.added[0][1][0]["file_name"] == "note.txt"
    assert store.added[0][1][0]["thread_id"] == "documents"

    results = asyncio.run(ingest_files(config, [path], store=store))
    assert results[0]["status"] == "duplicate"


def test_document_store_real_milvus_roundtrip(tmp_path):
    config = BotConfig(embed_dimensions=4)
    uri = str(tmp_path / "docs.db")
    store = DocumentStore(config, uri=uri, embedder=FakeEmbedder())
    path = tmp_path / "doc.txt"
    path.write_text("文档内容 " * 50, encoding="utf-8")

    try:
        results = asyncio.run(ingest_files(config, [path], store=store))
        assert results[0]["status"] == "imported"
        first_result = results[0]
        file_hash = hashlib.sha256(path.read_bytes()).hexdigest()
        assert asyncio.run(store.has_doc(file_hash))

        # 重复导入应被去重跳过
        results = asyncio.run(ingest_files(config, [path], store=store))
        assert results[0]["status"] == "duplicate"

        # 文档 collection 可被检索
        hits = asyncio.run(store.search_dense("文档", "", None, 5))
        assert hits and "文档内容" in hits[0]["content"]
        assert hits[0]["file_name"] == "doc.txt"

        # search_documents 工具端到端可用
        text_result = asyncio.run(search_documents("文档", store))
        assert "doc.txt" in text_result
        assert "文档内容" in text_result

        # 按 doc_id 删除整个文档后哈希去重失效
        deleted = asyncio.run(store.delete_doc(first_result["doc_id"]))
        assert deleted > 0
        assert not asyncio.run(store.has_doc(file_hash))
    finally:
        store.close()


def test_load_pdf_mineru_not_configured_returns_none(tmp_path, monkeypatch):
    """未配置 MinerU（无 endpoint 且无 API Key）时静默跳过，返回 None。"""
    from bot.package.knowledge import mineru_client

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    monkeypatch.setattr(mineru_client, "mineru_base_url", lambda config: None)
    config = BotConfig(_env_file=None, embed_dimensions=4)

    assert _load_pdf_mineru(path, config) is None


def test_load_pdf_mineru_returns_markdown_document(tmp_path, monkeypatch):
    """MinerU 在线返回 Markdown 时，包装为 markdown 格式的 Document。"""
    from bot.package.knowledge import mineru_client

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    markdown = "# 标题\n\n文档正文"

    monkeypatch.setattr(mineru_client, "mineru_base_url", lambda config: "https://mineru.net")
    monkeypatch.setattr(mineru_client, "parse_pdf", lambda path, config: markdown)
    config = BotConfig(_env_file=None, embed_dimensions=4, document_mineru_api_key="sk-abc")

    docs = _load_pdf_mineru(path, config)
    assert docs is not None
    assert len(docs) == 1
    assert docs[0].page_content == markdown
    assert docs[0].metadata["parser"] == "mineru"
    assert docs[0].metadata["format"] == "markdown"


def test_load_pdf_mineru_fallback_on_empty(tmp_path, monkeypatch):
    """MinerU 解析失败/返回空文本时返回 None，交由 LangChain/pypdf 降级。"""
    from bot.package.knowledge import mineru_client

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")

    monkeypatch.setattr(mineru_client, "mineru_base_url", lambda config: "https://mineru.net")
    monkeypatch.setattr(mineru_client, "parse_pdf", lambda path, config: "")
    config = BotConfig(_env_file=None, embed_dimensions=4, document_mineru_api_key="sk-abc")

    assert _load_pdf_mineru(path, config) is None



def test_load_pdf_mineru_agent_disabled_returns_none(tmp_path, monkeypatch):
    """Agent 轻量解析开关关闭时直接跳过。"""
    from bot.package.knowledge import mineru_client

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    called = {"agent": False}
    monkeypatch.setattr(
        mineru_client,
        "parse_pdf_agent",
        lambda path, config: called.update(agent=True) or "mocked",
    )
    config = BotConfig(
        _env_file=None,
        embed_dimensions=4,
        document_mineru_agent_enabled=False,
    )
    assert _load_pdf_mineru_agent(path, config) is None
    assert not called["agent"]


def test_load_pdf_mineru_agent_returns_markdown_document(tmp_path, monkeypatch):
    """Agent 轻量解析返回 Markdown 时包装为 Document。"""
    from bot.package.knowledge import mineru_client

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    markdown = "# Agent 标题\n\n正文"

    monkeypatch.setattr(mineru_client, "parse_pdf_agent", lambda path, config: markdown)
    config = BotConfig(_env_file=None, embed_dimensions=4)

    docs = _load_pdf_mineru_agent(path, config)
    assert docs is not None
    assert docs[0].page_content == markdown
    assert docs[0].metadata["parser"] == "mineru-agent"
    assert docs[0].metadata["format"] == "markdown"


def test_load_pdf_mineru_agent_fallback_on_empty(tmp_path, monkeypatch):
    """Agent 返回空文本时返回 None，继续下一重降级。"""
    from bot.package.knowledge import mineru_client

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")

    monkeypatch.setattr(mineru_client, "parse_pdf_agent", lambda path, config: "")
    config = BotConfig(_env_file=None, embed_dimensions=4)

    assert _load_pdf_mineru_agent(path, config) is None


def test_load_pdf_three_tier_fallback_order(tmp_path, monkeypatch):
    """PDF 解析严格按 精准 → Agent → 本地 三级降级。"""
    from bot.package.knowledge import document_ingestion as di

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    calls: list[str] = []

    def fake_mineru(path, config):
        calls.append("precise")

    def fake_agent(path, config):
        calls.append("agent")
        return [Document(
            page_content="# agent result",
            metadata={"parser": "mineru-agent", "format": "markdown"},
        )]

    def fake_local(path):
        calls.append("local")
        return [Document(
            page_content="local text",
            metadata={"source": str(path), "page": 1},
        )]

    monkeypatch.setattr(di, "_load_pdf_mineru", fake_mineru)
    monkeypatch.setattr(di, "_load_pdf_mineru_agent", fake_agent)
    monkeypatch.setattr(di, "_load_pdf_langchain", fake_local)
    config = BotConfig(_env_file=None, embed_dimensions=4)

    docs = di._load_pdf(path, config)
    assert [d.metadata.get("parser") for d in docs] == ["mineru-agent"]
    # 精准失败后才试 Agent，命中后不再走本地
    assert calls == ["precise", "agent"]


def test_load_pdf_three_tier_full_fallback_to_local(tmp_path, monkeypatch):
    """精准与 Agent 都失败时落到本地 LangChain/pypdf。"""
    from bot.package.knowledge import document_ingestion as di

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    calls: list[str] = []

    monkeypatch.setattr(
        di, "_load_pdf_mineru",
        lambda path, config: calls.append("precise") or None,
    )
    monkeypatch.setattr(
        di, "_load_pdf_mineru_agent",
        lambda path, config: calls.append("agent") or None,
    )
    monkeypatch.setattr(
        di, "_load_pdf_langchain",
        lambda path: calls.append("local") or [Document(
            page_content="local text",
            metadata={"source": str(path), "page": 1, "parser": "langchain", "format": "text"},
        )],
    )
    config = BotConfig(_env_file=None, embed_dimensions=4)

    docs = di._load_pdf(path, config)
    assert [d.metadata.get("parser") for d in docs] == ["langchain"]
    assert calls == ["precise", "agent", "local"]

